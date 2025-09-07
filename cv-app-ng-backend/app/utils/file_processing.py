"""
File processing utilities for document extraction.
"""
import fitz
import docx
from ..utils.debug import print_step

def extract_text_from_pdf(file_stream: bytes) -> str:
    """
    Extract text from PDF file stream.
    
    Args:
        file_stream: PDF file as bytes
        
    Returns:
        Extracted text content
    """
    print_step("PDF Text Extraction", {"file_size": len(file_stream)}, "input")
    doc = fitz.open(stream=file_stream, filetype="pdf")
    text = "".join(page.get_text() for page in doc)
    print_step("PDF Text Extraction", {
        "extracted_text_length": len(text), 
        "page_count": len(doc)
    }, "output")
    return text

def extract_text_from_docx(file_stream: bytes) -> str:
    """
    Extract text from DOCX file stream.
    
    Args:
        file_stream: DOCX file as bytes
        
    Returns:
        Extracted text content
    """
    print_step("DOCX Text Extraction", {"file_size": len(file_stream)}, "input")
    doc = docx.Document(file_stream)
    text = "\n".join([para.text for para in doc.paragraphs])
    print_step("DOCX Text Extraction", {
        "extracted_text_length": len(text), 
        "paragraph_count": len(doc.paragraphs)
    }, "output")
    return text
