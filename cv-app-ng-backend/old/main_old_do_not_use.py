# main.py
try:
    import torch
    TORCH_AVAILABLE = True
except ImportError:
    TORCH_AVAILABLE = False
    print("WARNING: PyTorch not available. Using CPU for ML models.")

from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from pydantic import BaseModel
import openai
from openai import AsyncOpenAI
import os
from dotenv import load_dotenv
import fitz
import docx
import base64
import json
import asyncio
import numpy as np
import jinja2
from weasyprint import HTML

# --- Production-Grade Imports ---
from langchain_pinecone import PineconeVectorStore
from langchain_chroma import Chroma
from langchain_openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
# from sentence_transformers import SentenceTransformer, CrossEncoder # Commented out - requires torch
from pinecone import Pinecone as PineconeClient, ServerlessSpec

# Conditional imports for ragas (has compatibility issues with uvloop)
try:
    from ragas import evaluate
    from ragas.metrics import faithfulness, answer_relevancy, context_precision, context_recall
    from datasets import Dataset
    RAGAS_AVAILABLE = True
    print("INFO:     Ragas evaluation framework loaded successfully.")
except ImportError as e:
    print(f"WARNING:  Ragas not available: {e}")
    RAGAS_AVAILABLE = False
    # Create dummy functions to prevent errors
    def evaluate(*args, **kwargs):
        return None
    def faithfulness(*args, **kwargs):
        return None
    def answer_relevancy(*args, **kwargs):
        return None
    def context_precision(*args, **kwargs):
        return None
    def context_recall(*args, **kwargs):
        return None
    class Dataset:
        @staticmethod
        def from_dict(data):
            return None

# --- Debug Flags ---
DEBUG = True
VERBOSE = True

def print_step(step_name, data=None, data_type="info"):
    """Helper function to print formatted debug information"""
    if DEBUG:
        print(f"\n{'='*60}")
        print(f"🔍 STEP: {step_name}")
        print(f"{'='*60}")
        if data is not None:
            if data_type == "input":
                print(f"📥 INPUT DATA:")
            elif data_type == "output":
                print(f"📤 OUTPUT DATA:")
            elif data_type == "error":
                print(f"❌ ERROR:")
            else:
                print(f"ℹ️  DATA:")
            
            if isinstance(data, (dict, list)):
                print(json.dumps(data, indent=2, default=str))
            else:
                print(str(data))
        print(f"{'='*60}\n")

# --- Environment and Model Loading ---
print_step("Environment Setup", "Loading environment variables from .env file")
load_dotenv()

# Initialize OpenAI client with error handling
print_step("OpenAI Client Initialization", {"api_key_present": bool(os.getenv("OPENAI_API_KEY"))}, "input")
try:
    client_async = AsyncOpenAI(
        api_key=os.getenv("OPENAI_API_KEY"),
        base_url="https://api.openai.com/v1"
    )
    print_step("OpenAI Client Initialization", "OpenAI client initialized successfully", "output")
except Exception as e:
    print_step("OpenAI Client Initialization", str(e), "error")
    print("INFO:     Some features may not work without a valid OpenAI API key.")
    client_async = None

# --- NEW: GPU/CPU Detection Logic ---
print_step("Device Detection", {"torch_available": TORCH_AVAILABLE}, "input")
if TORCH_AVAILABLE:
    device = 'cuda' if torch.cuda.is_available() else 'cpu'
    print_step("Device Detection", {"device": device, "cuda_available": torch.cuda.is_available()}, "output")
else:
    device = 'cpu'
    print_step("Device Detection", {"device": device, "reason": "PyTorch not available"}, "output")
# ------------------------------------

# --- Mocking Logic for Pinecone ---
MOCK_PINECONE = os.getenv("MOCK_PINECONE", "true").lower() == "true"
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
PINECONE_INDEX_NAME = "cv-architect-index"

print_step("Pinecone Configuration", {
    "mock_pinecone": MOCK_PINECONE,
    "api_key_present": bool(PINECONE_API_KEY),
    "index_name": PINECONE_INDEX_NAME
}, "input")

# --- OpenAI Setup ---
# Client already initialized above with error handling

# --- LangChain & Ragas Setup ---
print_step("Embeddings Initialization", {"client_available": client_async is not None}, "input")
if client_async is not None:
    embeddings = OpenAIEmbeddings(model="text-embedding-ada-002")
    print_step("Embeddings Initialization", "OpenAI embeddings initialized successfully", "output")
else:
    embeddings = None
    print_step("Embeddings Initialization", "OpenAI embeddings not available - API key required", "error")
    
print_step("Text Splitter Setup", {"chunk_size": 1000, "chunk_overlap": 200}, "input")
text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
print_step("Text Splitter Setup", "Text splitter initialized", "output")

vectorstore = None

print_step("Vectorstore Initialization", {"embeddings_available": embeddings is not None, "mock_pinecone": MOCK_PINECONE}, "input")
if embeddings is not None:
    if MOCK_PINECONE:
        print_step("Vectorstore Initialization", "Using mocked Pinecone (ChromaDB in-memory)", "info")
        vectorstore = Chroma(embedding_function=embeddings, collection_name=PINECONE_INDEX_NAME)
        print_step("Vectorstore Initialization", "ChromaDB vectorstore created", "output")
    else:
        print_step("Vectorstore Initialization", "Connecting to production Pinecone", "info")
        pinecone_client = PineconeClient(api_key=PINECONE_API_KEY)
        existing_indexes = pinecone_client.list_indexes().names()
        print_step("Pinecone Index Check", {"existing_indexes": existing_indexes}, "info")
        
        if PINECONE_INDEX_NAME not in existing_indexes:
            print_step("Pinecone Index Creation", {"index_name": PINECONE_INDEX_NAME, "dimension": 1536}, "input")
            pinecone_client.create_index(
                name=PINECONE_INDEX_NAME,
                dimension=1536,
                metric='cosine',
                spec=ServerlessSpec(cloud='aws', region='us-east-1')
            )
            print_step("Pinecone Index Creation", f"Index '{PINECONE_INDEX_NAME}' created successfully", "output")
        
        # Use the new langchain-pinecone package
        vectorstore = PineconeVectorStore.from_existing_index(
            index_name=PINECONE_INDEX_NAME,
            embedding=embeddings
        )
        print_step("Vectorstore Initialization", "Pinecone vectorstore connected", "output")
else:
    print_step("Vectorstore Initialization", "Vectorstore not initialized - OpenAI API key required", "error")

# --- Jinja2 Template Environment Setup ---
print_step("Jinja2 Template Setup", "Initializing template environment", "input")
template_loader = jinja2.FileSystemLoader(searchpath="./templates")
template_env = jinja2.Environment(loader=template_loader)
print_step("Jinja2 Template Setup", "Template environment initialized", "output")

# --- FastAPI App Initialization & CORS ---
print_step("FastAPI App Initialization", "Creating FastAPI application", "input")
app = FastAPI()
origins = ["http://localhost:5173", "http://127.0.0.1:5173"]
print_step("CORS Configuration", {"origins": origins}, "input")
app.add_middleware(CORSMiddleware, allow_origins=origins, allow_credentials=True, allow_methods=["*"], allow_headers=["*"])
print_step("FastAPI App Initialization", "FastAPI app and CORS middleware configured", "output")

# --- Data Models & Helpers ---
class CVRequest(BaseModel):
    job_description: str
    user_cv_text: str

class ImageRequest(BaseModel):
    image_base_64: str

class EvaluationRequest(BaseModel):
    job_description: str
    cv_json: dict # Accepts the JSON object of the generated CV

# --- CV Templating Models ---
class PersonalInfo(BaseModel):
    name: str = "Your Name"
    email: str = "your.email@example.com"
    phone: str = "+1234567890"
    location: str = "City, State"
    website: str = "your-website.com"
    linkedin: str = "linkedin.com/in/username"
    github: str = "github.com/username"

class Experience(BaseModel):
    company: str
    role: str
    startDate: str
    endDate: str
    location: str
    description: str
    achievements: list[str] = []

class Education(BaseModel):
    institution: str
    degree: str
    field: str
    startDate: str
    endDate: str
    gpa: str = ""

class Project(BaseModel):
    name: str
    description: str
    tech_stack: list[str] = []
    link: str = ""

class Skills(BaseModel):
    technical: list[str] = []
    soft: list[str] = []
    languages: list[str] = []

class CVData(BaseModel):
    personal: PersonalInfo
    professional_summary: str = ""
    experience: list[Experience] = []
    education: list[Education] = []
    projects: list[Project] = []
    skills: Skills
    licenses_certifications: list[dict] = []

class PDFRequest(BaseModel):
    templateId: str
    data: CVData

def extract_text_from_pdf(file_stream):
    print_step("PDF Text Extraction", {"file_size": len(file_stream)}, "input")
    doc = fitz.open(stream=file_stream, filetype="pdf")
    text = "".join(page.get_text() for page in doc)
    print_step("PDF Text Extraction", {"extracted_text_length": len(text), "page_count": len(doc)}, "output")
    return text

def extract_text_from_docx(file_stream):
    print_step("DOCX Text Extraction", {"file_size": len(file_stream)}, "input")
    doc = docx.Document(file_stream)
    text = "\n".join([para.text for para in doc.paragraphs])
    print_step("DOCX Text Extraction", {"extracted_text_length": len(text), "paragraph_count": len(doc.paragraphs)}, "output")
    return text

async def evaluate_with_persona(persona: str, job_description: str, cv_content: str):
    print_step("Persona Evaluation", {
        "persona": persona,
        "job_description_length": len(job_description),
        "cv_content_length": len(cv_content)
    }, "input")
    
    prompt = f"""
    You will act as: {persona}.
    Your task is to score the provided CV based on the job description from this perspective.
    Return JSON with "persona", "score", "justification".
    IMPORTANT: The "persona" field in your JSON response must exactly match the role you are acting as: "{persona}". Do not use any other name or value for this field.
    JOB: {job_description}
    CV: {cv_content}
    """

    
    
    print_step("OpenAI API Call", {"model": "gpt-3.5-turbo", "prompt_length": len(prompt)}, "input")
    response = await client_async.chat.completions.create(
        model="gpt-3.5-turbo", messages=[{"role": "user", "content": prompt}],
        temperature=0.0, response_format={"type": "json_object"}
    )
    
    result = json.loads(response.choices[0].message.content)
    print_step("Persona Evaluation", result, "output")
    return result

# --- API Endpoints ---
@app.get("/")
def read_root():
    return {"status": "CV Generator API is online"}

@app.post("/tailor-cv")
async def tailor_cv(request: CVRequest):
    print_step("CV Tailoring Request", {
        "job_description_length": len(request.job_description),
        "user_cv_text_length": len(request.user_cv_text)
    }, "input")

    print_step("Document Creation", {"cv_text_length": len(request.user_cv_text)}, "input")
    docs = text_splitter.create_documents([request.user_cv_text])
    print_step("Document Creation", {"document_count": len(docs), "total_chunks": sum(len(doc.page_content) for doc in docs)}, "output")
    
    print_step("Vectorstore Cleanup", {"mock_pinecone": MOCK_PINECONE}, "input")
    if MOCK_PINECONE:
        collection_ids = vectorstore.get()['ids']
        if collection_ids:
            vectorstore._collection.delete(ids=collection_ids)
            print_step("Vectorstore Cleanup", {"deleted_ids_count": len(collection_ids)}, "output")
        else:
            print_step("Vectorstore Cleanup", "No existing documents to delete", "output")
    
    print_step("Document Indexing", {"document_count": len(docs)}, "input")
    vectorstore.add_documents(docs)
    print_step("Document Indexing", "Documents added to vectorstore", "output")

    print_step("Document Retrieval", {"job_description": request.job_description, "k": 7}, "input")
    retriever = vectorstore.as_retriever(search_kwargs={'k': 7})
    retrieved_docs = retriever.invoke(request.job_description)
    retrieved_context = "\n\n".join([doc.page_content for doc in retrieved_docs])
    print_step("Document Retrieval", {
        "retrieved_docs_count": len(retrieved_docs),
        "retrieved_context_length": len(retrieved_context),
        "retrieved_context_preview": retrieved_context[:200] + "..." if len(retrieved_context) > 200 else retrieved_context
    }, "output")
    
    generation_prompt = f"""
    You are a meticulous AI CV Architect with expertise in creating compelling, ATS-optimized resumes. Your task is to analyze the provided CV content and job description, then generate a tailored CV that maximizes the candidate's chances of landing the specific role.

    INSTRUCTIONS:
    1. Extract relevant information from the RAW CV CONTEXT
    2. Identify key requirements from the TARGET JOB DESCRIPTION
    3. Create a tailored CV that highlights the most relevant experience, skills, and achievements
    4. Use action verbs and quantify achievements where possible
    5. Ensure the CV is ATS-friendly with clear section headers
    6. Maintain professional formatting and structure

    DATA:
    - RAW CV CONTEXT: --- {retrieved_context} ---
    - TARGET JOB DESCRIPTION: --- {request.job_description} ---

    Return a JSON object with the following structure:
    {{
        "name": "Full Name",
        "contact": {{
            "email": "email@example.com",
            "phone": "+1234567890",
            "location": "City, State",
            "linkedin": "linkedin.com/in/username",
            "portfolio": "portfolio-url.com"
        }},
        "professional_summary": "2-3 sentence summary tailored to the job",
        "professional_experience": [
            {{
                "company": "Company Name",
                "position": "Job Title",
                "duration": "Start Date - End Date",
                "location": "City, State",
                "achievements": [
                    "Quantified achievement 1",
                    "Quantified achievement 2",
                    "Quantified achievement 3"
                ]
            }}
        ],
        "education": [
            {{
                "institution": "University Name",
                "degree": "Degree Type",
                "field": "Field of Study",
                "graduation_year": "YYYY",
                "gpa": "X.XX (if relevant)"
            }}
        ],
        "licenses_certifications": [
            {{
                "name": "Certification Name",
                "issuer": "Issuing Organization",
                "date": "YYYY",
                "expiry": "YYYY (if applicable)"
            }}
        ],
        "skills": {{
            "technical": ["Skill 1", "Skill 2", "Skill 3"],
            "soft": ["Skill 1", "Skill 2", "Skill 3"],
            "languages": ["Language 1", "Language 2"]
        }}
    }}

    Provide ONLY the raw JSON object as your response.
    """
    
    print_step("CV Generation", {
        "model": "gpt-4o",
        "prompt_length": len(generation_prompt),
        "temperature": 0.1
    }, "input")
    
    try:
        response = await client_async.chat.completions.create(
            model="gpt-4o", messages=[{"role": "user", "content": generation_prompt}],
            temperature=0.1, response_format={"type": "json_object"}
        )
        print_step("OpenAI Response", {
            "response_length": len(response.choices[0].message.content),
            "model_used": response.model,
            "tokens_used": response.usage.total_tokens if response.usage else "unknown"
        }, "output")
        
        structured_content = json.loads(response.choices[0].message.content)
        print_step("Structured Content Parsing", "JSON successfully parsed", "output")
        
        # Debug: Show the actual generated content
        print_step("Generated CV Content Preview", {
            "name": structured_content.get("name", "NOT_FOUND"),
            "contact_email": structured_content.get("contact", {}).get("email", "NOT_FOUND") if structured_content.get("contact") else "NO_CONTACT_SECTION",
            "summary_length": len(structured_content.get("professional_summary", "")),
            "experience_count": len(structured_content.get("professional_experience", [])),
            "education_count": len(structured_content.get("education", [])),
            "skills_technical_count": len(structured_content.get("skills", {}).get("technical", [])) if structured_content.get("skills") else 0
        }, "output")

        # Ragas evaluation (if available)
        print_step("RAGAS Evaluation Setup", {"ragas_available": RAGAS_AVAILABLE}, "input")
        if RAGAS_AVAILABLE:
            try:
                print_step("RAGAS Dataset Creation", {
                    "question_length": len(request.job_description),
                    "contexts_count": len(retrieved_docs),
                    "answer_length": len(json.dumps(structured_content))
                }, "input")
                
                # Validate that we have contexts to evaluate
                if not retrieved_docs:
                    print_step("RAGAS Dataset Creation", "No retrieved documents - skipping RAGAS evaluation", "error")
                    ragas_scores = {
                        "faithfulness": 0.0,
                        "answer_relevancy": 0.0,
                        "context_precision": 0.0,
                        "context_recall": 0.0
                    }
                else:
                    # Create a proper reference answer for context_precision metric
                    reference_answer = f"Based on the job description: {request.job_description}, the CV should highlight relevant skills and experience."
                    
                    # Fix: contexts should be a list of strings, not nested list
                    # Fix: use ground_truths (plural) not ground_truth
                    dataset = Dataset.from_dict({
                        'question': [request.job_description],
                        'contexts': [doc.page_content for doc in retrieved_docs],  # List of strings, not nested
                        'answer': [json.dumps(structured_content)],
                        'ground_truths': [reference_answer]  # Use ground_truths (plural)
                    })
                    print_step("RAGAS Dataset Creation", "Dataset created successfully", "output")
                    
                    print_step("RAGAS Evaluation Execution", {
                        "metrics": ["faithfulness", "answer_relevancy", "context_precision", "context_recall"]
                    }, "input")
                    ragas_result = await asyncio.to_thread(evaluate, dataset, metrics=[faithfulness, answer_relevancy, context_precision, context_recall])
                    ragas_scores = ragas_result.to_pandas().to_dict('records')[0]
                    print_step("RAGAS Evaluation Execution", "Evaluation completed", "output")
                    
                    # Handle NaN values in RAGAS scores
                    print_step("RAGAS Score Processing", ragas_scores, "input")
                    for key, value in ragas_scores.items():
                        if isinstance(value, float) and (np.isnan(value) or np.isinf(value)):
                            ragas_scores[key] = 0.0
                    print_step("RAGAS Score Processing", "NaN values handled", "output")
                        
            except Exception as e:
                print_step("RAGAS Evaluation", str(e), "error")
                ragas_scores = {
                    "faithfulness": 0.0,
                    "answer_relevancy": 0.0,
                    "context_precision": 0.0,
                    "context_recall": 0.0
                }
        else:
            print_step("RAGAS Evaluation", "RAGAS not available, using default scores", "info")
            ragas_scores = {
                "faithfulness": 0.0,
                "answer_relevancy": 0.0,
                "context_precision": 0.0,
                "context_recall": 0.0
            }

        print_step("Committee Evaluation Setup", {
            "personas": ["Strict Hiring Manager", "Creative Recruiter", "Senior Technical Lead"],
            "cv_content_length": len(json.dumps(structured_content))
        }, "input")
        
        personas = ["Strict Hiring Manager", "Creative Recruiter", "Senior Technical Lead"]
        evaluation_tasks = [evaluate_with_persona(p, request.job_description, json.dumps(structured_content)) for p in personas]
        
        print_step("Committee Evaluation Execution", {"task_count": len(evaluation_tasks)}, "input")
        committee_evaluations = await asyncio.gather(*evaluation_tasks)
        print_step("Committee Evaluation Execution", {"completed_evaluations": len(committee_evaluations)}, "output")
        
        # Handle potential NaN values in committee scores
        print_step("Committee Score Processing", committee_evaluations, "input")
        scores = []
        for e in committee_evaluations:
            score = e.get('score', 0)
            if isinstance(score, (int, float)) and not (np.isnan(score) or np.isinf(score)):
                scores.append(score)
            else:
                scores.append(0)
        
        committee_analysis = {
            "individual_evaluations": committee_evaluations,
            "average_score": round(np.mean(scores), 2) if scores else 0.0
        }
        print_step("Committee Score Processing", committee_analysis, "output")
        
        print_step("Final Analysis Assembly", {
            "ragas_scores": ragas_scores,
            "committee_analysis": committee_analysis
        }, "input")
        
        structured_content['analysis'] = {
            "ragas_scores": ragas_scores,
            "committee_evaluation": committee_analysis
        }
        
        print_step("CV Tailoring Complete", {
            "final_content_keys": list(structured_content.keys()),
            "analysis_present": 'analysis' in structured_content
        }, "output")
        
        return structured_content

    except Exception as e:
        print_step("CV Tailoring Error", str(e), "error")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/tailor-cv-from-file/")
async def tailor_cv_from_file(job_description: str, cv_file: UploadFile = File(...)):
    print_step("File Upload Request", {
        "filename": cv_file.filename,
        "content_type": cv_file.content_type,
        "job_description_length": len(job_description)
    }, "input")
    
    file_content = await cv_file.read()
    print_step("File Reading", {"file_size": len(file_content)}, "output")
    
    print_step("File Type Detection", {"filename": cv_file.filename}, "input")
    if cv_file.filename.endswith(".pdf"):
        user_cv_text = extract_text_from_pdf(file_content)
    elif cv_file.filename.endswith(".docx"):
        user_cv_text = extract_text_from_docx(file_content)
    else:
        print_step("File Type Detection", "Unsupported file type", "error")
        raise HTTPException(status_code=400, detail="Unsupported file type.")
    
    print_step("CV Request Creation", {
        "job_description_length": len(job_description),
        "extracted_text_length": len(user_cv_text)
    }, "input")
    rag_request = CVRequest(job_description=job_description, user_cv_text=user_cv_text)
    print_step("CV Request Creation", "Request object created, calling tailor_cv", "output")
    
    return await tailor_cv(rag_request)

@app.post("/extract-cv-data")
async def extract_cv_data(request: dict):
    """
    Extract structured CV data from raw CV text using AI.
    """
    print_step("CV Data Extraction Request", {
        "cv_text_length": len(request.get("cv_text", "")),
        "job_description_length": len(request.get("job_description", ""))
    }, "input")

    try:
        cv_text = request.get("cv_text", "")
        job_description = request.get("job_description", "")
        
        if not cv_text:
            raise HTTPException(status_code=400, detail="CV text is required")
        
        # Use the AI service to extract structured CV data
        structured_content = await ai_service.extract_cv_data(cv_text, job_description)
        
        print_step("CV Data Extraction", {
            "extracted_fields": list(structured_content.keys()) if isinstance(structured_content, dict) else "Unknown"
        }, "output")
        
        return structured_content
        
    except Exception as e:
        print_step("CV Data Extraction Error", str(e), "error")
        raise HTTPException(status_code=500, detail=f"Error extracting CV data: {e}")

@app.post("/transcribe-audio/")
async def transcribe_audio(audio_file: UploadFile = File(...)):
    print_step("Audio Transcription Request", {
        "filename": audio_file.filename,
        "content_type": audio_file.content_type
    }, "input")
    
    file_location = f"/tmp/{audio_file.filename}"
    try:
        print_step("Audio File Processing", {"file_location": file_location}, "input")
        with open(file_location, "wb+") as file_object:
            file_object.write(audio_file.file.read())
        print_step("Audio File Processing", "File written to temporary location", "output")
        
        print_step("Whisper Transcription", {"model": "whisper-1"}, "input")
        with open(file_location, "rb") as audio:
            transcription = await client_async.audio.transcriptions.create(model="whisper-1", file=audio)
        print_step("Whisper Transcription", {
            "transcription_length": len(transcription.text),
            "transcription_preview": transcription.text[:100] + "..." if len(transcription.text) > 100 else transcription.text
        }, "output")
        
        return {"transcription": transcription.text}
    finally:
        if os.path.exists(file_location):
            print_step("Cleanup", "Removing temporary audio file", "info")
            os.remove(file_location)

@app.post("/analyze-jd-image/")
async def analyze_jd_image(request: ImageRequest):
    print_step("Image Analysis Request", {
        "image_base64_length": len(request.image_base_64)
    }, "input")
    
    try:
        print_step("GPT-4 Vision Analysis", {
            "model": "gpt-4o",
            "max_tokens": 500
        }, "input")
        
        response = await client_async.chat.completions.create(
          model="gpt-4o",
          messages=[{
              "role": "user", "content": [
                {"type": "text", "text": "Analyze this image of a job description. Extract the key responsibilities and required skills. Return the result as a clean block of text."},
                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{request.image_base_64}"}}
              ]}], max_tokens=500,
        )
        
        extracted_text = response.choices[0].message.content
        print_step("GPT-4 Vision Analysis", {
            "extracted_text_length": len(extracted_text),
            "extracted_text_preview": extracted_text[:200] + "..." if len(extracted_text) > 200 else extracted_text
        }, "output")
        
        return {"extracted_job_description": extracted_text}
    except Exception as e:
        print_step("Image Analysis Error", str(e), "error")
        raise HTTPException(status_code=500, detail=f"Error analyzing image: {e}")

# --- API Endpoints ---

# ... (keep the / endpoint)

@app.post("/evaluate-cv")
async def evaluate_cv(request: EvaluationRequest):
    """
    Performs a committee evaluation on a provided CV JSON against a job description.
    """
    print_step("Committee Evaluation Request", {
        "job_description_length": len(request.job_description),
        "cv_keys": list(request.cv_json.keys())
    }, "input")

    try:
        # Convert the CV JSON object back to a string for the LLM prompt
        cv_content_str = json.dumps(request.cv_json, indent=2)

        print_step("Committee Evaluation Setup", {
            "personas": ["Strict Hiring Manager", "Creative Recruiter", "Senior Technical Lead"],
            "cv_content_length": len(cv_content_str)
        }, "input")
        
        personas = ["Strict Hiring Manager", "Creative Recruiter", "Senior Technical Lead"]
        evaluation_tasks = [evaluate_with_persona(p, request.job_description, cv_content_str) for p in personas]
        
        print_step("Committee Evaluation Execution", {"task_count": len(evaluation_tasks)}, "input")
        committee_evaluations = await asyncio.gather(*evaluation_tasks)
        print_step("Committee Evaluation Execution", {"completed_evaluations": len(committee_evaluations)}, "output")
        
        # Handle potential invalid score values
        print_step("Committee Score Processing", committee_evaluations, "input")
        scores = []
        for e in committee_evaluations:
            score = e.get('score', 0)
            if isinstance(score, (int, float)) and not (np.isnan(score) or np.isinf(score)):
                scores.append(score)
            else:
                scores.append(0) # Default to 0 if score is invalid or missing
        
        committee_analysis = {
            "individual_evaluations": committee_evaluations,
            "average_score": round(np.mean(scores), 2) if scores else 0.0
        }
        print_step("Committee Score Processing", committee_analysis, "output")
        
        return committee_analysis

    except Exception as e:
        print_step("Committee Evaluation Error", str(e), "error")
        raise HTTPException(status_code=500, detail=f"Error during evaluation: {e}")

@app.post("/generate-pdf")
async def generate_pdf(request: PDFRequest):
    """
    Generate a PDF from CV data using the specified template.
    """
    print_step("PDF Generation Request", {
        "template_id": request.templateId,
        "personal_name": request.data.personal.name,
        "experience_count": len(request.data.experience),
        "education_count": len(request.data.education)
    }, "input")
    
    try:
        # Check if template exists
        template_file = f"{request.templateId}.html"
        template_path = os.path.join("./templates", template_file)
        
        if not os.path.exists(template_path):
            print_step("PDF Generation Error", f"Template not found: {template_file}", "error")
            raise HTTPException(status_code=404, detail=f"Template '{request.templateId}' not found")
        
        print_step("Template Loading", {"template_file": template_file}, "input")
        template = template_env.get_template(template_file)
        print_step("Template Loading", "Template loaded successfully", "output")
        
        # Render the HTML with the user's data
        print_step("HTML Rendering", {"data_keys": list(request.data.model_dump().keys())}, "input")
        html_content = template.render(request.data.model_dump())
        print_step("HTML Rendering", {"html_length": len(html_content)}, "output")
        
        # Generate PDF from the rendered HTML
        print_step("PDF Generation", {"html_length": len(html_content)}, "input")
        pdf_bytes = HTML(string=html_content).write_pdf()
        print_step("PDF Generation", {"pdf_size_bytes": len(pdf_bytes)}, "output")
        
        # Set headers for file download
        headers = {
            'Content-Disposition': f'attachment; filename="cv_{request.data.personal.name.replace(" ", "_")}.pdf"'
        }
        
        print_step("PDF Generation Complete", {
            "pdf_size_kb": round(len(pdf_bytes) / 1024, 2),
            "filename": f"cv_{request.data.personal.name.replace(' ', '_')}.pdf"
        }, "output")
        
        return Response(pdf_bytes, headers=headers, media_type='application/pdf')
        
    except Exception as e:
        print_step("PDF Generation Error", str(e), "error")
        raise HTTPException(status_code=500, detail=f"Error generating PDF: {e}")


# ... (keep the rest of the endpoints: /tailor-cv, /tailor-cv-from-file, etc.)

# --- Application Startup Message ---
print_step("Application Startup", "CV Generator API is ready to serve requests!", "output")
print("\n" + "="*80)
print("🚀 CV GENERATOR API STARTED SUCCESSFULLY")
print("="*80)
print("📋 Available Endpoints:")
print("   • GET  /                    - Health check")
print("   • POST /tailor-cv           - Tailor CV from text")
print("   • POST /tailor-cv-from-file - Tailor CV from uploaded file")
print("   • POST /evaluate-cv         - Perform committee evaluation on a generated CV")
print("   • POST /generate-pdf        - Generate PDF from CV data using templates")
print("   • POST /transcribe-audio    - Transcribe audio to text")
print("   • POST /analyze-jd-image    - Extract job description from image")
print("="*80)
print("🔧 Debug Mode: ENABLED - Detailed logging will be shown for each request")
print("="*80 + "\n")
